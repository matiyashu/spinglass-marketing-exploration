"""Build the Word manual from docs/DASHBOARD_MANUAL.md + captured screenshots.

Renders the markdown (headings, tables, lists, blockquotes, inline bold/italic/
code/links) into a .docx and embeds the dashboard screenshots after their
matching section heading.

Run (after capture_dashboard_screens.py):
    python backend/scripts/build_manual_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
MD = ROOT / "docs" / "DASHBOARD_MANUAL.md"
IMG = ROOT / "assets" / "manual"
OUT = ROOT / "docs" / "Spin-Glass_Dashboard_Manual.docx"

TEAL = RGBColor(0x0D, 0x94, 0x88)
INK = RGBColor(0x1F, 0x2A, 0x37)
MUTED = RGBColor(0x64, 0x74, 0x8B)

# Section-number → (image file, caption). Matched against "### N.N ..." headings.
IMAGES = {
    "4.1": ("02_home.png", "Workspace Home — KPI tiles and portfolio coherence for the active brand."),
    "4.2": ("03_data_setup.png", "Data setup — the five tables and what each one unlocks."),
    "4.4": ("04_method_status.png", "Method status — which analyses the supplied data enables."),
    "5.1": ("05_portfolio.png", "Portfolio overview — brands ranked by memory synchronisation."),
    "5.2": ("06_memory_map.png", "Brand memory map — the 10×10 coupling matrix with reinforcing (teal) and conflicting (red) cells."),
    "5.3": ("07_tensions.png", "Brand tensions — strongest negative couplings and frustrated triads."),
    "5.4": ("08_competitive_leakage.png", "Competitive leakage — competitor-pattern overlap over time."),
    "6.1": ("09_vertical_overview.png", "Vertical overview — per-product coherence grouped by vertical."),
    "6.2": ("10_product_memory_fit.png", "Product memory fit — memory synchronisation by product."),
    "6.3": ("11_segment_differences.png", "Segment differences — cross-segment memory overlap matrix."),
    "6.4": ("12_switching_risk.png", "Switching risk — products ranked by competitor overlap."),
    "7.1": ("17_campaign_overview.png", "Campaign overview — campaigns and their outcome-data status."),
    "7.2": ("18_creative_memory.png", "Creative memory pattern — the declared target vector."),
    "7.3": ("19_field_response.png", "Field response — per-association movement during the flight (waterfall)."),
    "7.4": ("20_simulator.png", "Scenario simulator — simulated retrieval vs spend (no purchase-probability card)."),
    "7.5": ("21_validation.png", "Observed validation — outcome series when outcome data exists."),
    "8.1": ("13_rolling_regime.png", "Rolling regime — coherence indicators wave over wave, campaign bands shaded."),
    "8.2": ("15_persistence.png", "Memory persistence — post-campaign residue by campaign."),
    "8.3": ("14_replicas.png", "Replica / fragmentation — estimation-stability vs landscape P(q), kept separate."),
    "8.4": ("16_stress_test.png", "Stress test — association susceptibility (campaign sensitivity)."),
}


def shade(cell, hexcolor: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\)")


def add_runs(paragraph, text: str, base_size: int = 10) -> None:
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()]).font.size = Pt(base_size)
        bold, ital, code, link, _url = m.groups()
        if bold is not None:
            r = paragraph.add_run(bold); r.bold = True
        elif ital is not None:
            r = paragraph.add_run(ital); r.italic = True
        elif code is not None:
            r = paragraph.add_run(code); r.font.name = "Consolas"; r.font.color.rgb = TEAL
        else:
            r = paragraph.add_run(link); r.font.color.rgb = TEAL
        r.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:]).font.size = Pt(base_size)


def add_image(doc: Document, fname: str, caption: str) -> None:
    path = IMG / fname
    if not path.exists():
        return
    w, h = Image.open(path).size
    ratio = h / w
    target_w = 6.6
    if target_w * ratio > 7.7:
        target_w = 7.7 / ratio
    doc.add_picture(str(path), width=Inches(target_w))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = MUTED


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, cellp in enumerate(row):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], cellp, base_size=8.5)
            if ri == 0:
                shade(cell, "0D9488")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def main() -> None:
    lines = MD.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Base style + narrow margins.
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.7)

    # --- Cover ---
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Spin-Glass\nBrand Memory & Campaign Dynamics"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = INK
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Program Manual — features, formulas, how to read results, and use cases"); r.font.size = Pt(13); r.font.color.rgb = MUTED
    doc.add_paragraph()
    add_image(doc, "01_landing.png", "")
    v = doc.add_paragraph(); v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = v.add_run("V3 · marketing-first dashboard"); r.font.size = Pt(10); r.font.color.rgb = TEAL
    doc.add_page_break()

    i = 0
    seen_title = False
    while i < len(lines):
        line = lines[i].rstrip()

        # Tables: a run of pipe rows.
        if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").replace(" ", "")) <= set("-:"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = []
            for bi, brow in enumerate(block):
                if bi == 1:
                    continue  # separator
                cells = [c.strip() for c in brow.strip().strip("|").split("|")]
                rows.append(cells)
            add_table(doc, rows)
            continue

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            i += 1
            if not seen_title:
                seen_title = True  # cover already has the title
                continue
            h = doc.add_heading(line[2:], level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            text = line[4:]
            doc.add_heading(text, level=2)
            m = re.match(r"(\d+\.\d+)", text)
            if m and m.group(1) in IMAGES:
                fname, cap = IMAGES[m.group(1)]
                add_image(doc, fname, cap)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, line[2:], base_size=10)
            for run in p.runs:
                run.italic = True
            i += 1
            continue
        if re.match(r"^\s*[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*] ", "", line))
            i += 1
            continue
        if re.match(r"^\s*\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\s*\d+\. ", "", line))
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(str(OUT))
    print(f"saved -> {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
